from flask import Flask, render_template, request, jsonify, redirect, session, url_for
from flask_cors import CORS
from training.web_crawler import crawl_website
import uuid
import sqlite3
import os
import google.generativeai as genai
from agents.agent_manager import create_agent
from training.pdf_loader import load_pdf
from training.url_loader import load_url
from training.vector_store import store_data
from training.vector_store import search_data
from chatbot.chatbot_engine import ask_agent
from config import GEMINI_API_KEY

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = "data/uploaded_pdfs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

app.secret_key = "secret123"


# Gemini API Key
genai.configure(api_key=GEMINI_API_KEY)

model = genai.GenerativeModel("gemini-2.5-flash")

# DATABASE CONNECTION
def get_db():
    conn = sqlite3.connect("users.db")
    conn.row_factory = sqlite3.Row
    return conn


# LOGIN PAGE
@app.route("/")
def login_page():
    return render_template("login.html")


# LOGIN
@app.route("/login", methods=["GET","POST"])
def login():

    if request.method == "POST":

        username = request.form.get("username").strip()
        password = request.form.get("password").strip()

        conn = get_db()
        cur = conn.cursor()

        # fetch user first
        cur.execute("SELECT * FROM users WHERE username=?", (username,))
        user = cur.fetchone()

        conn.close()

        if user:
            if user["password"] == password:
                session["user"] = username
                return redirect("/index")
            else:
                return "Incorrect password"
        else:
            return "User not found"

    return render_template("login.html")


# REGISTER
@app.route("/register", methods=["GET", "POST"])
def register():

    if request.method == "POST":

        username = request.form["username"]
        password = request.form["password"]

        conn = get_db()
        cur = conn.cursor()

        cur.execute(
            "INSERT INTO users (username,password) VALUES (?,?)",
            (username, password)
        )

        conn.commit()
        conn.close()

        return redirect("/login")

    return render_template("register.html")


# HOME PAGE
@app.route("/index")
def home():

    if "user" not in session:
        return redirect("/login")

    return render_template("index.html")


# LOGOUT
@app.route("/logout")
def logout():

    session.pop("user", None)

    return redirect("/login")


# get agent information when user log in again
@app.route("/get_agent_data", methods=["GET"])
def get_agent_data():

    if "user" not in session:
        return jsonify({"agent": None, "training": []})

    username = session["user"]

    conn = get_db()
    cur = conn.cursor()

    cur.execute("SELECT * FROM agents WHERE username=?", (username,))
    agent = cur.fetchone()

    if not agent:
        conn.close()
        return jsonify({"agent": None, "training": []})

    agent_id = agent["id"]

    cur.execute("""
        SELECT id, file_name, source_type, source_value 
        FROM training_data WHERE agent_id=?
    """, (agent_id,))

    training = cur.fetchall()

    conn.close()

    return jsonify({
        "agent": dict(agent),
        "training": [dict(row) for row in training]
    })

# delete
@app.route("/delete_training/<int:id>", methods=["DELETE"])
def delete_training(id):

    conn = get_db()
    cur = conn.cursor()

    cur.execute("DELETE FROM training_data WHERE id=?", (id,))

    conn.commit()
    conn.close()

    return jsonify({"status": "deleted"})


# UPDATE PROMPT
@app.route("/update_prompt", methods=["POST"])
def update_prompt():

    data = request.json
    agent_id = data.get("agent_id")
    prompt = data.get("prompt")

    conn = get_db()
    cur = conn.cursor()

    cur.execute("""
        UPDATE agents SET prompt=? WHERE id=?
    """, (prompt, agent_id))

    conn.commit()
    conn.close()

    return jsonify({"status": "updated"})

# Build Agent
@app.route("/build", methods=["POST"])
def build():

    if "user" not in session:
        return jsonify({"error": "Login required"})

    username = session["user"]

    data = request.json
    prompt = data.get("prompt")
    name = data.get("name")

    conn = get_db()
    cur = conn.cursor()

    #  CHECK if user already has agent
    cur.execute("SELECT * FROM agents WHERE username=?", (username,))
    existing = cur.fetchone()

    if existing:
        conn.close()
        return jsonify({
            "error": "You already created an agent. You can update it."
        })

    #  CREATE NEW AGENT
    agent_id = str(uuid.uuid4())

    cur.execute("""
        INSERT INTO agents (id, username, name, prompt)
        VALUES (?, ?, ?, ?)
    """, (agent_id, username, name, prompt))

    conn.commit()
    conn.close()

    return jsonify({"agent_id": agent_id})


# TRAIN AGENT

@app.route("/train", methods=["POST"])
def train():

    if "user" not in session:
        return jsonify({
            "error": "Login required"
        }), 401

    username = session["user"]

    conn = get_db()
    cur = conn.cursor()

    # -------------------------
    # GET AGENT
    # -------------------------
    cur.execute(
        "SELECT id FROM agents WHERE username=?",
        (username,)
    )

    agent = cur.fetchone()

    if not agent:
        conn.close()
        return jsonify({
            "error": "Create agent first"
        }), 400

    agent_id = agent["id"]

    text = ""
    files_processed = 0

    print("\n==============================")
    print("TRAINING STARTED")
    print("USER :", username)
    print("AGENT :", agent_id)
    print("==============================")

    # ===================================================
    # FILE TRAINING
    # ===================================================

    files = request.files.getlist("files")

    print("FILES RECEIVED :", len(files))

    for file in files:

        if not file or file.filename.strip() == "":
            continue

        try:

            unique_name = (
                str(uuid.uuid4()) + "_" + file.filename
            )

            filepath = os.path.join(
                UPLOAD_FOLDER,
                unique_name
            )

            file.save(filepath)

            print("Saved :", filepath)

            ext = file.filename.split(".")[-1].lower()

            extracted = ""

            if ext == "pdf":

                extracted = load_pdf(filepath)

            elif ext in ["doc", "docx"]:

                from docx import Document

                doc = Document(filepath)

                extracted = "\n".join(
                    p.text for p in doc.paragraphs
                )

            elif ext == "txt":

                with open(
                    filepath,
                    "r",
                    encoding="utf-8"
                ) as f:

                    extracted = f.read()

            elif ext in ["jpg", "jpeg", "png"]:

                extracted = ""

                print("Image uploaded (OCR disabled)")

            else:

                print("Unsupported file :", ext)

            if extracted.strip():

                text += "\n" + extracted

                print(
                    "Extracted characters :",
                    len(extracted)
                )

            cur.execute(
                """
                INSERT INTO training_data
                (
                    agent_id,
                    file_name,
                    file_path,
                    source_type,
                    source_value
                )
                VALUES (?,?,?,?,?)
                """,
                (
                    agent_id,
                    file.filename,
                    filepath,
                    "file",
                    file.filename
                )
            )

            files_processed += 1

        except Exception as e:

            print("FILE ERROR :", e)

    # ===================================================
    # WEBSITE CRAWLER
    # ===================================================

    crawl_url = request.form.get("crawl_url")

    if crawl_url:

        crawl_url = crawl_url.strip()

    if crawl_url:

        try:

            print("\n==============================")
            print("CRAWLER START")
            print(crawl_url)
            print("==============================")

            extracted = crawl_website(
                crawl_url,
                max_pages=20
            )

            print("Crawler Finished")

            print(
                "Crawler Text Length :",
                len(extracted)
            )

            if extracted and extracted.strip():

                text += "\n" + extracted

                print("Crawler Success")

            else:

                print("No content extracted from crawler")

            cur.execute(
                """
                INSERT INTO training_data
                (
                    agent_id,
                    file_name,
                    file_path,
                    source_type,
                    source_value
                )
                VALUES (?,?,?,?,?)
                """,
                (
                    agent_id,
                    None,
                    None,
                    "crawler",
                    crawl_url
                )
            )

        except Exception as e:

            import traceback

            print("\n========== CRAWLER ERROR ==========")

            traceback.print_exc()

    # ===================================================
    # FINAL TRAINING
    # ===================================================

    text = text.strip()

    print("\n==============================")
    print("TOTAL TEXT LENGTH :", len(text))
    print("==============================")

    if len(text) == 0:

        conn.close()

        return jsonify({

            "error": "No content extracted"

        }), 400

    try:

        print("\nVECTOR TRAINING STARTED")

        text = (
            text
            .replace("\n", " ")
            .replace("\t", " ")
        )

        store_data(
            text=text,
            agent_id=agent_id
        )

        conn.commit()

        print("VECTOR STORE COMPLETED")

        conn.close()

        return jsonify({

            "status": "trained",

            "files_processed": files_processed,

            "message": "Training completed successfully"

        })

    except Exception as e:

        import traceback

        print("\n========== STORE DATA ERROR ==========")

        traceback.print_exc()

        conn.rollback()

        conn.close()

        return jsonify({

            "error": str(e)

        }), 500

# @app.route("/train", methods=["POST"])
# def train():

#     if "user" not in session:
#         return jsonify({"error": "Login required"}), 401

#     username = session["user"]

#     conn = get_db()
#     cur = conn.cursor()

#     #  Get user's agent
#     cur.execute("SELECT id FROM agents WHERE username=?", (username,))
#     agent = cur.fetchone()

#     if not agent:
#         conn.close()
#         return jsonify({"error": "Create agent first"})

#     agent_id = agent["id"]

#     text = ""
#     files_processed = 0

    
#     #  MULTIPLE FILES
    
#     files = request.files.getlist("files")

#     print("FILES RECEIVED:", files)

#     for file in files:

#         if not file or file.filename.strip() == "":
#             continue

#         try:
#             # unique file name
#             unique_name = str(uuid.uuid4()) + "_" + file.filename
#             filepath = os.path.join(UPLOAD_FOLDER, unique_name)

#             file.save(filepath)
#             print("Saved file:", filepath)

#             ext = file.filename.split(".")[-1].lower()

           
#             #  TEXT EXTRACTION
            
#             if ext == "pdf":
#                 extracted = load_pdf(filepath)

#             elif ext in ["doc", "docx"]:
#                 from docx import Document
#                 doc = Document(filepath)
#                 extracted = "\n".join([p.text for p in doc.paragraphs])

#             elif ext in ["jpg", "jpeg", "png"]:
#                 extracted = ""  # no OCR for now
#                 print("Image uploaded:", filepath)

#             elif ext == "txt":
#                 with open(filepath, "r", encoding="utf-8") as f:
#                     extracted = f.read()

#             else:
#                 extracted = ""
#                 print("Unsupported file type:", ext)

#             text += extracted

          
#             #  STORE IN DATABASE
           
#             cur.execute("""
#                 INSERT INTO training_data 
#                 (agent_id, file_name, file_path, source_type, source_value)
#                 VALUES (?, ?, ?, ?, ?)
#             """, (
#                 agent_id,
#                 file.filename,
#                 filepath,
#                 "file",
#                 file.filename
#             ))

#             files_processed += 1

#         except Exception as e:
#             print("❌ File error:", e)

    
    
#     #  CRAWLER TRAINING

#     print("===== Train Start =====")
#     crawl_url = request.form.get("crawl_url")
#     print("URL RECEIVED:", crawl_url)

#     if crawl_url and crawl_url.strip() != "":
#         try:
#             print("CRAWLER:", crawl_url)
#             print("Calling Crawler...")
#             extracted = crawl_website(crawl_url, max_pages=20)
#             print("Crwal Finished")
#             print("Extracted Length:", len(extracted))
#             print(extracted[:500])
#             print("TEXT LENGTH:", len(extracted))  # 🔥 DEBUG
#             if extracted.strip():
#                 text += extracted
#             else:
#                 print("⚠️ No content from crawler")

#             cur.execute("""
#                 INSERT INTO training_data 
#                 (agent_id, file_name, file_path, source_type, source_value)
#                 VALUES (?, ?, ?, ?, ?)
#             """, (agent_id, None, None, "crawler", crawl_url))

        
#         except Exception as e:
#             print("❌ Crawler error:", e)
        



# 🔹 VECTOR STORE (FINAL FIX)
# =========================

    if text.strip():

        print("🔥 TRAINING STARTED")
        print("TEXT LENGTH:", len(text))

        # ✅ CLEAN TEXT (IMPORTANT)
        text = text.replace("\n", " ").replace("\t", " ")

        # ✅ STORE ONCE ONLY
        # store_data(text, agent_id)
        print("🔥 CALLING STORE DATA...")
        store_data(text, agent_id)
        print("🔥 STORE DATA FINISHED")

        print("✅ TRAINING COMPLETED")

        conn.commit()
        conn.close()

        return jsonify({
            "status": "trained",
            "files_processed": files_processed
        })

    else:
        conn.commit()
        conn.close()

        return jsonify({
            "error": "No content extracted"
        }), 400
    

# chat
@app.route("/chat", methods=["POST"])
def chat():

    data = request.json
    message = data.get("message")
    agent_id = data.get("agent_id")

    message_lower = message.lower()

    # ✅ GREETING HANDLING (ADD THIS)
    greetings = ["hello", "hi", "hey", "good morning", "good evening", "good afternoon"]

    if any(greet in message_lower for greet in greetings):
        return jsonify({
            "response": "Hello 👋 How can I assist you today?"
        })

    # 🔎 SEARCH
    context, distance = search_data(message, agent_id)

    print("\n====================")
    print("QUESTION:", message)
    print("DISTANCE:", distance)
    print("CONTEXT:\n", context[:1000])
    print("====================\n")

    # ❌ NO DATA FOUND
    if not context or distance > 10:
        return jsonify({
            "response": "I could not find this information in the provided data."
        })

    # ✅ STRICT PROMPT
    final_prompt = f"""
        You are a strict AI assistant.

        RULES:
        - Answer ONLY from CONTEXT
        - Do NOT use outside knowledge
        - Do NOT guess
        - If answer is explicitly stated, return it.
        - If answer is implied from context, infer only from context.
        - Do not invent technologies not mentioned.
        "I could not find this information in the provided data."


        RESPONSE FORMAT (MANDATORY):
        - Always give response in bullet points.
        - Each point must start with •
        - Give a short meaningful description for each point.
        - Keep each point 1-2 lines only.
        - Include important details only.
        - Keep response clean, professional and well formatted.
        - Do NOT return one-word answers.
        - Do NOT return long paragraphs.

        Example Format:

        • Website Development:
        Builds responsive and modern websites for businesses.

        • Mobile Application:
        Develops user-friendly mobile solutions for various needs.

        CONTEXT:
        {context}

        QUESTION:
        {message}
        """

    try:
        response = model.generate_content(final_prompt)
        answer = response.text.strip()

        # 🔥 HARD FILTER
        bad_words = ["typically", "generally", "usually", "co-founder is", "in general"]

        if any(word in answer.lower() for word in bad_words):
            answer = "I could not find this information in the provided data."

    except Exception as e:
        print("GEMINI ERROR:", e)
        answer = "⚠️ Daily limit reached.Please try again later."

    return jsonify({"response": answer})

# @app.route("/chat", methods=["POST"])
# def chat():

#     data = request.json
#     message = data.get("message")
#     agent_id = data.get("agent_id")

#     conn = get_db()
#     cur = conn.cursor()

#     cur.execute("SELECT prompt FROM agents WHERE id=?", (agent_id,))
#     agent = cur.fetchone()

#     conn.close()

#     if not agent:
#         return jsonify({"response": "Agent not found"})

#     message_lower = message.lower()

#     # ✅ Greeting shortcut
#     greetings = ["hello", "hi", "hey", "good morning", "good evening", "good afternoon"]
#     if any(word in message_lower for word in greetings):
#         return jsonify({
#             "response": "Hello 👋 How can I help you today?"
#         })

#     # 🔎 SEARCH FROM VECTOR DB
#     context, distance = search_data(message, agent_id)

#     # 🔥 DEBUG (VERY IMPORTANT)
#     print("\n====================")
#     print("QUESTION:", message)
#     print("DISTANCE:", distance)
#     print("CONTEXT:\n", context[:1000])
#     print("====================\n")

#     # ❌ HARD STOP (NO CONTEXT → NO ANSWER)
#     if not context or distance > 5:
#         return jsonify({
#             "response": "I could not find this information in the provided data."
#         })

#     # ✅ STRICT PROMPT (NO ESCAPE)
#     final_prompt = f"""
# You are a strict AI assistant.

# RULES:
# - Answer ONLY from the given CONTEXT
# - Do NOT use your own knowledge
# - Do NOT guess
# - Do NOT explain generally
# - If exact answer is not present → reply EXACTLY:
# "I could not find this information in the provided data."

# FORMAT:
# - Use bullet points
# - Keep answers short and clear

# CONTEXT:
# {context}

# QUESTION:
# {message}
# """

#     try:
#         response = model.generate_content(final_prompt)
#         answer = response.text.strip()

#         # 🔥 EXTRA SAFETY FILTER (VERY IMPORTANT)
#         if "co-founder is someone" in answer.lower() or "typically" in answer.lower():
#             answer = "I could not find this information in the provided data."

#     except Exception as e:
#         print("GEMINI ERROR:", e)
#         answer = "⚠️ Daily limit reached. Please try again later."

#     return jsonify({"response": answer})





# EMBED CODE PAGE
@app.route("/embed/<agent_id>")
def embed(agent_id):

    return render_template(
        "embed_code.html",
        agent_id=agent_id
    )


if __name__ == "__main__":
    # app.run(debug=True)
    app.run(host="0.0.0.0", port=5000, debug=True)


# # Chat
# @app.route("/chat", methods=["POST"])
# def chat():

#     data = request.json
#     message = data.get("message")
#     agent_id = data.get("agent_id")

#     conn = get_db()
#     cur = conn.cursor()

#     cur.execute("SELECT prompt FROM agents WHERE id=?", (agent_id,))
#     agent = cur.fetchone()

#     conn.close()

#     if not agent:
#         return jsonify({"response": "Agent not found"})

#     message_lower = message.lower()

#     # Greeting detection
#     greetings = ["hello", "hi", "hey", "good morning", "good evening", "good afternoon"]

#     if any(word in message_lower for word in greetings):
#         return jsonify({
#             "response": "Hello 👋 How can I help you today?"
#         })

#     # agent domain
#     agent_prompt = agent["prompt"].lower()
#     domain = agent_prompt.replace("create", "").replace("agent", "").strip()

#     # 🔎 Search training data
#     context, distance = search_data(message, agent_id)
#     print("DISTANCE:", distance)
#     print("CONTEXT:", context[:500])

#     print("FINAL CONTEXT:\n", context)


#     # Decide prompt
#     # if context.strip() != "" and distance < 10:
#     if context and distance < 3.0:
#         final_prompt = f"""
# You are an expert assistant in {domain}.

# STRICT RULES:
# - Answer ONLY from the provided training data
# - DO NOT guess or generate your own information
# - If answer is not clearly present → reply:
#   "I could not find this information in the provided data.


# FORMAT:
# - Use bullet points
# - Keep answer clean and professional
# - Do not mention "training data"

# Training Data:
# {context}

# Question:
# {message}
# """

#     else:
        
#         final_prompt = f"""
# You are an expert assistant in the field of {domain}.

# Rules:
# 1. First determine if the user question belongs to the domain: {agent_prompt}

# 2. If the question is related to the domain:
#    - If training data contains the answer → answer using it.


# IMPORTANT RULES:
# 1. If answer is present → answer clearly
# 2. If partially present → try to infer
# 3. DO NOT say Not available in data

# If the question is unrelated to {domain}, reply exactly:
# "I am a {domain} agent and I can only answer questions related to {domain}."

# Question:
# {message}
# """

#     try:

#         response = model.generate_content(final_prompt)
#         answer = response.text

#     except Exception as e:
#         print("GEMINI ERROR:", e)
#         answer = "⚠️ Daily limit reached. Please try again later."

#     return jsonify({"response": answer})











# from flask import Flask, render_template, request, jsonify, redirect, session, url_for
# from flask_cors import CORS
# import sqlite3


# from agents.agent_manager import create_agent
# from training.pdf_loader import load_pdf
# from training.url_loader import load_url
# from training.vector_store import store_data
# from chatbot.chatbot_engine import ask_agent

# app = Flask(__name__)
# CORS(app)  
# app.secret_key = "secret123"

# # Database
# def get_db():
#     return sqlite3.connect("users.db")

# # Login Page
# @app.route("/")
# def login_page():
#     return render_template ("login.html")

# # LOGIN
# @app.route("/login", methods=["GET","POST"])
# def login():

#     if request.method == "POST":

#         username = request.form.get("username")
#         password = request.form.get("password")

#         conn = sqlite3.connect("users.db")
#         cur = conn.cursor()

#         cur.execute(
#             "SELECT * FROM users WHERE username=? AND password=?",
#             (username, password)
#         )

#         user = cur.fetchone()

#         conn.close()

#         if user:
#             return redirect(url_for("home"))   # 👈 redirect to index page
#         else:
#             return "Invalid username or password"

#     return render_template("login.html")
    
# # Register
# @app.route("/register", methods = ['GET', 'POST'])
# def register():
#     if request.method == "POST":
#         username = request.form["username"]
#         password = request.form["password"]

#         conn = sqlite3.connect("users.db")
#         cur = conn.cursor()
#         cur.execute("INSERT INTO users VALUES (?,?)", (username, password))
#         conn.commit()
#         conn.close()
#         return redirect("/login")
#     return render_template("register.html")

# # HOME
# @app.route("/index")
# def home():
#     # if "user" not in session:
#     #     return redirect("/login")
#     return render_template("index.html")


# # LOGOUT
# @app.route("/logout")
# def logout():
#     session.pop("user", None)
#     return redirect("/login")


# # BUILD TAB
# @app.route("/build", methods=["POST"])
# def build():

#     data = request.json

#     prompt = data["prompt"]
#     name = data["name"]

#     agent_id = create_agent(prompt, name)

#     return jsonify({
#         "agent_id": agent_id
#     })
# # @app.route("/build", methods=["POST"])
# # def build():

# #     data = request.json

# #     prompt = data["prompt"]
# #     name = data["name"]

# #     agent_id = create_agent(prompt, name)

# #     return jsonify({"agent_id": agent_id})



# # TRAIN TAB
# @app.route("/train", methods=["POST"])
# def train():

#     if "file" in request.files:
#         file = request.files["file"]
#         text = load_pdf(file)

#     if "url" in request.form:
#         url = request.form["url"]
#         text = load_url(url)

#     store_data(text)

#     return jsonify({"status": "trained"})

# # CHAT
# @app.route("/chat", methods=["POST"])
# def chat():

#     data = request.json
#     message = data.get("message")
#     agent_id = data.get("agent_id")

#     response = f"Agent {agent_id} says: {message}"

#     return jsonify({"response": response})

# # @app.route("/chat", methods=["POST"])
# # def chat():
# #     question = request.json["message"]
# #     response = ask_agent(question)
# #     return jsonify({"response": response})

# # PUBLISH
# @app.route("/embed/<agent_id>")
# def embed(agent_id):
#     return render_template("embed_code.html", agent_id=agent_id)

# if __name__ == "__main__":
#     app.run(debug=True)